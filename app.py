import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta, date
import json
import os
import pickle
import numpy as np

CONTIFY_LOGO = "https://www.contify.com/wp-content/uploads/2026/02/contify-logo.svg"
st.set_page_config(page_title="Contify SEO Weekly Review", page_icon=CONTIFY_LOGO,
                   layout="wide", initial_sidebar_state="expanded")

# =====================================================
# PATHS & CONFIG
# =====================================================
APP_DIR = os.path.dirname(os.path.abspath(__file__))
TOKEN_FILE = os.path.join(APP_DIR, "token.pickle")
CONFIG_FILE = os.path.join(APP_DIR, "config.json")
SERVICE_ACCOUNT_FILE = os.path.join(os.path.expanduser("~"), "Downloads", "seo-report-weekly-054644ce3e32.json")
OAUTH_CLIENT_FILE = os.path.join(os.path.expanduser("~"), "Downloads", "client_secret_903915771943-pngfdafiam32olbubqq8kiqb60lh9f73.apps.googleusercontent.com.json")

SCOPES = ['https://www.googleapis.com/auth/analytics.readonly', 'https://www.googleapis.com/auth/webmasters.readonly']
DEFAULT_CONFIG = {"ga4_property": "250811000", "gsc_site": "https://www.contify.com/", "auth_method": "demo",
                  "targets": {"overall_traffic": 5538, "organic_traffic": 1154, "overall_users": 4562, "organic_users": 706, "mqls": 13}}

AI_SOURCES = ['chatgpt.com', 'chat.openai.com', 'gemini.google.com', 'claude.ai', 'perplexity.ai',
              'copilot.microsoft.com', 'you.com', 'bard.google.com', 'poe.com', 'phind.com']
EUROPE_COUNTRIES = ['Austria', 'Finland', 'France', 'Germany', 'Ireland', 'Italy', 'Netherlands', 'Sweden']
EXCLUDE_FROM_OTHER = ['India', 'Indonesia', 'Turkey']

# =====================================================
# CSS — Contify Brand, White BG, Clean & Spacious
# =====================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
@import url('https://fonts.googleapis.com/icon?family=Material+Icons');
@import url('https://fonts.googleapis.com/css2?family=Material+Symbols+Rounded:opsz,wght,FILL,GRAD@20..48,100..700,0..1,-50..200&display=block');
@import url('https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined:opsz,wght,FILL,GRAD@20..48,100..700,0..1,-50..200&display=block');

/* Force Material Symbols font on Streamlit icon buttons (fixes ligature text) */
[data-testid="stSidebarCollapseButton"], [data-testid="stSidebarCollapseButton"] *,
[data-testid="stSidebarCollapsedControl"], [data-testid="stSidebarCollapsedControl"] *,
button[kind="headerNoPadding"], button[kind="headerNoPadding"] *,
[data-testid="baseButton-headerNoPadding"], [data-testid="baseButton-headerNoPadding"] * {
  font-family: 'Material Symbols Rounded', 'Material Icons' !important;
  font-weight: normal !important;
  font-style: normal !important;
  font-size: 24px !important;
  line-height: 1 !important;
  letter-spacing: normal !important;
  text-transform: none !important;
  white-space: nowrap !important;
  word-wrap: normal !important;
  direction: ltr !important;
  -webkit-font-feature-settings: 'liga' !important;
  font-feature-settings: 'liga' !important;
  -webkit-font-smoothing: antialiased !important;
  text-rendering: optimizeLegibility !important;
}

/* === Force Light Theme & Global Reset === */
html, body, [data-testid="stApp"], [data-testid="stApp"] > div { background-color: #ffffff !important; color: #1e293b !important; }
html, body, [class*="css"], [data-testid="stApp"] * { font-family: 'Inter', sans-serif; }
.main .block-container { max-width: 1200px; padding: 2rem 2rem; }

/* === Sidebar — White & Clean === */
[data-testid="stSidebar"] { background-color: #ffffff !important; border-right: 1px solid #e5e7eb !important; }
[data-testid="stSidebar"] * { color: #1e293b !important; }
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] h2,
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] strong { color: #0f3460 !important; }
[data-testid="stSidebar"] hr { border-color: #e5e7eb !important; }

/* === Streamlit Widget Overrides === */
[data-testid="stNumberInput"] input,
[data-testid="stTextInput"] input { background: #f8fafc !important; border: 1px solid #d1d5db !important; border-radius: 8px !important; color: #1e293b !important; }
[data-testid="stNumberInput"] input:focus,
[data-testid="stTextInput"] input:focus { border-color: #1a56db !important; box-shadow: 0 0 0 2px rgba(26,86,219,0.15) !important; }
/* Number input +/- stepper buttons — blend into the input field as one cohesive control */
[data-testid="stNumberInput"] { display: flex; align-items: stretch; }
[data-testid="stNumberInput"] > div:first-child { flex: 1; }
[data-testid="stNumberInput"] button,
[data-testid="stNumberInputContainer"] button {
    background: #f8fafc !important;
    color: #0f3460 !important;
    border: 1px solid #d1d5db !important;
    border-radius: 0 !important;
    margin: 0 !important;
    min-width: 32px !important;
    box-shadow: none !important;
    transition: background 0.15s ease !important;
}
[data-testid="stNumberInput"] button:first-of-type { border-radius: 0 !important; border-left: none !important; }
[data-testid="stNumberInput"] button:last-of-type { border-radius: 0 8px 8px 0 !important; border-left: none !important; }
[data-testid="stNumberInput"] button:hover { background: #eff6ff !important; color: #1a56db !important; }
[data-testid="stNumberInput"] button svg,
[data-testid="stNumberInputContainer"] button svg { fill: currentColor !important; stroke: currentColor !important; }
/* Make the input itself attached on the right side */
[data-testid="stNumberInput"] input { border-radius: 8px 0 0 8px !important; border-right: none !important; }
/* Selectbox / dropdown widget — force light surface */
[data-testid="stSelectbox"] > div > div,
[data-baseweb="select"] > div { background: #f8fafc !important; color: #1e293b !important; border-color: #d1d5db !important; }
[data-baseweb="select"] svg { fill: #0f3460 !important; }
[data-baseweb="popover"] [role="listbox"], [data-baseweb="menu"] { background: #ffffff !important; color: #1e293b !important; }
[data-baseweb="popover"] li, [data-baseweb="menu"] li { background: #ffffff !important; color: #1e293b !important; }
[data-baseweb="popover"] li:hover, [data-baseweb="menu"] li:hover { background: #eff6ff !important; }
/* Show/hide checkbox toggle (e.g., "Show password") */
[data-testid="stCheckbox"] > label > div:first-child { background: #f1f5f9 !important; border: 1px solid #d1d5db !important; }
button[kind="secondary"], [data-testid="stButton"] button {
    background: linear-gradient(135deg, #0f3460, #1a56db) !important; color: #ffffff !important;
    border: none !important; border-radius: 8px !important; font-weight: 600 !important; padding: 8px 20px !important;
    transition: opacity 0.2s !important;
}
button[kind="secondary"]:hover, [data-testid="stButton"] button:hover { opacity: 0.9 !important; }
[data-testid="stRadio"] label { color: #374151 !important; }
[data-testid="stRadio"] [data-testid="stMarkdownContainer"] p { font-weight: 500 !important; }

/* === Dashboard Header === */
.dashboard-header {
    background: linear-gradient(135deg, #0f3460 0%, #1a56db 100%);
    padding: 35px 40px; border-radius: 16px; color: white; margin-bottom: 30px;
}
.dashboard-header h1 { margin: 0; font-size: 1.9rem; font-weight: 800; letter-spacing: -0.5px; color: #ffffff !important; }
.dashboard-header p { margin: 8px 0 0; opacity: 0.85; font-size: 0.95rem; color: #ffffff !important; }

/* === Section Cards === */
.section-card {
    background: #ffffff; border: 1px solid #e2e8f0; border-radius: 14px;
    padding: 24px 28px; margin-bottom: 26px; box-shadow: 0 2px 8px rgba(15,52,96,0.04);
}
.section-title {
    font-size: 1.1rem; font-weight: 700; color: #ffffff; margin: -24px -28px 0 -28px;
    padding: 14px 28px; background: linear-gradient(135deg, #0f3460, #1a56db);
    border-radius: 14px 14px 0 0; display: flex; align-items: center; gap: 8px;
    letter-spacing: -0.2px; border-bottom: none;
}

/* === Insight Box === */
.insight-box {
    background: linear-gradient(135deg, #eff6ff, #f0f7ff); border-left: 5px solid #1a56db; padding: 18px 22px;
    border-radius: 0 12px 12px 0; margin: 18px 0 10px; font-size: 0.9rem; line-height: 1.8;
    border: 1px solid #dbeafe; border-left: 5px solid #1a56db;
}
.insight-box strong { color: #0f3460; }
.insight-box ul { margin: 10px 0 0; padding-left: 20px; }
.insight-box li { margin-bottom: 6px; }
.pos { color: #059669; font-weight: 700; }
.neg { color: #dc2626; font-weight: 700; }

/* === KPI Table === */
.kpi-table { width: 100%; border-collapse: separate; border-spacing: 0; border-radius: 12px; overflow: hidden; border: 1px solid #e2e8f0; }
.kpi-table th { background: #f0f4ff; color: #0f3460; padding: 12px 20px; font-size: 0.8rem; text-transform: uppercase; letter-spacing: 0.4px; font-weight: 700; text-align: left; border-bottom: 2px solid #dbe2ef; }
.kpi-table td { padding: 14px 20px; border-bottom: 1px solid #f0f2f5; font-size: 0.95rem; color: #1e293b; vertical-align: middle; }
.kpi-table tr:last-child td { border-bottom: none; }
.kpi-table tr:hover td { background: #f8faff; }

/* === Change Table (light header) === */
.change-table { width: 100%; border-collapse: separate; border-spacing: 0; border-radius: 12px; overflow: hidden; border: 1px solid #dbe2ef; font-size: 0.9rem; box-shadow: 0 2px 8px rgba(0,0,0,0.03); }
.change-table th { background: #f0f4ff; color: #0f3460; padding: 13px 18px; font-weight: 600; text-align: left; border-bottom: 2px solid #d1d5db; }
.change-table td { padding: 12px 18px; border-bottom: 1px solid #f0f2f5; color: #1e293b; }
.change-table tr:last-child td { border-bottom: none; }
.change-table tr:hover td { background: #f0f7ff; }
.change-table tr:nth-child(even) td { background: #fafbfc; }

/* === Clickable rows === */
.change-table tr.clickable-row { cursor: pointer; }
.change-table tr.clickable-row:hover td { background: #dbeafe !important; }

/* === Drilldown table === */
.drilldown-table { width: 100%; border-collapse: separate; border-spacing: 0; border-radius: 8px; overflow: hidden; border: 1px solid #dbeafe; font-size: 0.85rem; margin-top: 8px; }
.drilldown-table th { background: #dbeafe; color: #1e40af; padding: 10px 14px; font-weight: 600; text-align: left; }
.drilldown-table td { padding: 9px 14px; border-bottom: 1px solid #eff6ff; color: #1e293b; }
.drilldown-table tr:last-child td { border-bottom: none; }

/* === Metric Cards === */
.metric-mini { background: linear-gradient(145deg, #ffffff, #f8fafc); border: 1px solid #dbe2ef; border-radius: 14px; padding: 24px; text-align: center; box-shadow: 0 3px 10px rgba(15,52,96,0.06); transition: all 0.25s; }
.metric-mini:hover { box-shadow: 0 6px 20px rgba(15,52,96,0.12); transform: translateY(-2px); }
.metric-mini .val { font-size: 1.9rem; font-weight: 800; color: #0f3460; }
.metric-mini .label { font-size: 0.78rem; color: #6b7280; text-transform: uppercase; letter-spacing: 0.6px; margin-bottom: 8px; font-weight: 600; }
.metric-mini .change { font-size: 0.85rem; margin-top: 8px; font-weight: 600; }

/* === Tags === */
.tag-live { background: #10b981; color: white; padding: 3px 12px; border-radius: 20px; font-size: 0.7rem; font-weight: 700; }
.tag-demo { background: #f59e0b; color: white; padding: 3px 12px; border-radius: 20px; font-size: 0.7rem; font-weight: 700; }

/* === Better Week Labels === */
.week-current { background: linear-gradient(135deg, #1a56db, #3b82f6); color: white; padding: 4px 14px; border-radius: 8px; font-size: 0.8rem; font-weight: 700; display: inline-block; letter-spacing: 0.3px; }
.week-label-cell { font-size: 0.92rem; }
.week-label-cell strong { color: #0f3460; }

/* === Footer === */
.footer-text { text-align: center; color: #9ca3af !important; font-size: 0.85rem; padding: 20px 0 10px; border-top: 1px solid #e5e7eb; margin-top: 20px; }

/* === Plotly Chart Containers === */
[data-testid="stPlotlyChart"] {
    border-radius: 14px; overflow: hidden; border: 1px solid #e5e7eb;
    background: #fafbfc; padding: 8px; margin-bottom: 16px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.03);
}

/* === Expander clean up === */
[data-testid="stExpander"] { border: 1px solid #e5e7eb !important; border-radius: 10px !important; margin: 12px 0 !important; }
[data-testid="stExpander"] summary { font-weight: 600 !important; color: #1a56db !important; padding: 12px 16px !important; }
[data-testid="stExpander"] [data-testid="stExpanderDetails"] { padding: 0 16px 16px !important; }

/* === Hide Streamlit branding (keep header + collapsed-sidebar control visible) === */
#MainMenu, footer { visibility: hidden !important; }
header[data-testid="stHeader"] {
  background: transparent !important;
  box-shadow: none !important;
  height: 3.75rem !important;
  display: block !important;
  visibility: visible !important;
  opacity: 1 !important;
  z-index: 999990 !important;
  pointer-events: auto !important;
}
/* Hide only the right-side decoration/status, NEVER the toolbar (it contains the sidebar reopen button) */
header[data-testid="stHeader"] [data-testid="stDecoration"],
header[data-testid="stHeader"] [data-testid="stStatusWidget"] { display: none !important; }
/* Hide the actual menu/deploy items inside toolbar but keep toolbar itself visible */
[data-testid="stToolbar"] [data-testid="stMainMenu"],
[data-testid="stToolbar"] [data-testid="stToolbarActions"] { display: none !important; }
[data-testid="stToolbar"] { display: flex !important; visibility: visible !important; }

/* Force-show the reopen control across every Streamlit version / class name */
[data-testid="stExpandSidebarButton"],
[data-testid="stSidebarCollapsedControl"],
[data-testid="collapsedControl"],
button[data-testid="stExpandSidebarButton"],
button[aria-label="Open sidebar"],
button[kind="header"][aria-label*="ide"] {
  display: flex !important;
  visibility: visible !important;
  opacity: 1 !important;
  position: fixed !important;
  top: 12px !important;
  left: 12px !important;
  z-index: 999999 !important;
  background: linear-gradient(135deg, #0f3460, #1a56db) !important;
  border-radius: 10px !important;
  padding: 8px 12px !important;
  box-shadow: 0 4px 14px rgba(15,52,96,0.35) !important;
  pointer-events: auto !important;
  width: 44px !important;
  height: 44px !important;
  min-width: 44px !important;
  min-height: 44px !important;
  align-items: center !important;
  justify-content: center !important;
  cursor: pointer !important;
  transform: none !important;
}
[data-testid="stExpandSidebarButton"] *,
[data-testid="stSidebarCollapsedControl"] *,
[data-testid="collapsedControl"] *,
button[data-testid="stExpandSidebarButton"] *,
button[aria-label="Open sidebar"] * {
  color: #ffffff !important;
  fill: #ffffff !important;
  stroke: #ffffff !important;
  opacity: 1 !important;
  visibility: visible !important;
  width: auto !important;
  height: auto !important;
}
[data-testid="stExpandSidebarButton"] svg,
[data-testid="stSidebarCollapsedControl"] svg {
  width: 22px !important;
  height: 22px !important;
  fill: #ffffff !important;
}

/* Always-visible sidebar collapse button (inside sidebar) */
[data-testid="stSidebarCollapseButton"] {
  display: block !important;
  visibility: visible !important;
  opacity: 1 !important;
}
[data-testid="stSidebarCollapseButton"] button {
  opacity: 1 !important;
  visibility: visible !important;
  background: #eff6ff !important;
  color: #0f3460 !important;
  border-radius: 8px !important;
  border: 1px solid #dbeafe !important;
}

/* Always-visible reopen button (top-left, when sidebar is collapsed) */
[data-testid="stSidebarCollapsedControl"],
[data-testid="collapsedControl"] {
  display: flex !important;
  visibility: visible !important;
  opacity: 1 !important;
  position: fixed !important;
  top: 12px !important;
  left: 12px !important;
  z-index: 999999 !important;
  background: linear-gradient(135deg, #0f3460, #1a56db) !important;
  border-radius: 10px !important;
  padding: 4px 8px !important;
  box-shadow: 0 4px 14px rgba(15,52,96,0.35) !important;
  pointer-events: auto !important;
}
[data-testid="stSidebarCollapsedControl"] button,
[data-testid="collapsedControl"] button {
  opacity: 1 !important;
  visibility: visible !important;
  background: transparent !important;
  color: #ffffff !important;
  border: none !important;
}
</style>
""", unsafe_allow_html=True)

# =====================================================
# SESSION STATE
# =====================================================
if 'drilldown' not in st.session_state:
    st.session_state.drilldown = None

# =====================================================
# CONFIG HELPERS
# =====================================================
def _is_cloud():
    """Detect if running on Streamlit Cloud (no local config.json)"""
    return not os.path.exists(CONFIG_FILE)

def load_config():
    # On Streamlit Cloud, read from st.secrets
    if _is_cloud():
        cfg = DEFAULT_CONFIG.copy()
        cfg['ga4_property'] = st.secrets.get('ga4', {}).get('property_id', cfg['ga4_property'])
        cfg['gsc_site'] = st.secrets.get('ga4', {}).get('gsc_site', cfg['gsc_site'])
        cfg['auth_method'] = 'service_account'
        if 'targets' in st.secrets:
            cfg['targets'] = dict(st.secrets['targets'])
        if 'hubspot' in st.secrets:
            cfg['hubspot_token'] = st.secrets['hubspot']['token']
        return cfg
    # Locally, read from config.json
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r') as f:
            return {**DEFAULT_CONFIG, **json.load(f)}
    return DEFAULT_CONFIG.copy()

def save_config(c):
    if _is_cloud(): return  # Don't save on cloud
    with open(CONFIG_FILE, 'w') as f:
        json.dump(c, f, indent=2)

# =====================================================
# WEEK HELPERS (Mon–Sun)
# =====================================================
def get_last_n_weeks(n=5, ref_date=None):
    today = ref_date or datetime.now().date()
    last_sunday = today - timedelta(days=(today.weekday() + 1) % 7)
    if last_sunday == today:
        last_sunday -= timedelta(days=7)
    weeks = []
    for i in range(n):
        sun = last_sunday - timedelta(weeks=i)
        mon = sun - timedelta(days=6)
        weeks.append((mon, sun, f"{mon.strftime('%d %b')} – {sun.strftime('%d %b')}"))
    weeks.reverse()
    return weeks

def assign_weeks(df, weeks, date_col='Date', date_fmt='%Y%m%d'):
    df = df.copy()
    df['_date'] = pd.to_datetime(df[date_col], format=date_fmt, errors='coerce')
    df['Week'] = None; df['Week_Label'] = None; df['Week_Idx'] = -1
    for i, (m, s, lbl) in enumerate(weeks):
        mask = (df['_date'].dt.date >= m) & (df['_date'].dt.date <= s)
        df.loc[mask, 'Week'] = i; df.loc[mask, 'Week_Label'] = lbl; df.loc[mask, 'Week_Idx'] = i
    return df[df['Week'].notna()].copy()

# =====================================================
# AUTH
# =====================================================
def get_oauth_credentials():
    try:
        from google_auth_oauthlib.flow import InstalledAppFlow
        from google.auth.transport.requests import Request
    except ImportError:
        st.error("Missing: `pip install google-auth-oauthlib`"); return None
    creds = None
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as f: creds = pickle.load(f)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try: creds.refresh(Request())
            except Exception: creds = None
        if not creds:
            if not os.path.exists(OAUTH_CLIENT_FILE):
                st.error(f"OAuth file missing: `{OAUTH_CLIENT_FILE}`"); return None
            flow = InstalledAppFlow.from_client_secrets_file(OAUTH_CLIENT_FILE, SCOPES)
            for port in [8090, 8091, 8092, 8093, 8094]:
                try: creds = flow.run_local_server(port=port, prompt='consent', access_type='offline'); break
                except OSError: continue
        with open(TOKEN_FILE, 'wb') as f: pickle.dump(creds, f)
    return creds

def get_sa_credentials():
    from google.oauth2 import service_account
    # Try Streamlit Cloud secrets first
    try:
        if hasattr(st, 'secrets') and 'ga4' in st.secrets and 'credentials' in st.secrets['ga4']:
            sa = st.secrets["ga4"]["credentials"]
            # Explicitly cast each field to str to avoid Streamlit AttrDict serialization issues
            creds_info = {
                "type": str(sa.get("type", "service_account")),
                "project_id": str(sa["project_id"]),
                "private_key_id": str(sa["private_key_id"]),
                "private_key": str(sa["private_key"]).replace('\\n', '\n'),
                "client_email": str(sa["client_email"]),
                "client_id": str(sa["client_id"]),
                "auth_uri": str(sa.get("auth_uri", "https://accounts.google.com/o/oauth2/auth")),
                "token_uri": str(sa.get("token_uri", "https://oauth2.googleapis.com/token")),
                "auth_provider_x509_cert_url": str(sa.get("auth_provider_x509_cert_url", "https://www.googleapis.com/oauth2/v1/certs")),
                "client_x509_cert_url": str(sa.get("client_x509_cert_url", "")),
            }
            return service_account.Credentials.from_service_account_info(creds_info, scopes=SCOPES)
    except Exception as e:
        st.error(f"Cloud credentials error: {e}")
        return None
    # Fallback to local file
    if os.path.exists(SERVICE_ACCOUNT_FILE):
        return service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    return None

# =====================================================
# GA4 API FUNCTIONS
# =====================================================
def _ga4_client(creds):
    from google.analytics.data_v1beta import BetaAnalyticsDataClient
    return BetaAnalyticsDataClient(credentials=creds)

def _ga4_report(creds, prop, start, end, dims, metrics, dim_filter=None, limit=10000):
    from google.analytics.data_v1beta.types import RunReportRequest, DateRange, Dimension, Metric
    client = _ga4_client(creds)
    req_kwargs = dict(
        property=f"properties/{prop}",
        date_ranges=[DateRange(start_date=start, end_date=end)],
        dimensions=[Dimension(name=d) for d in dims],
        metrics=[Metric(name=m) for m in metrics],
        limit=limit,
    )
    if dim_filter:
        req_kwargs['dimension_filter'] = dim_filter
    return client.run_report(RunReportRequest(**req_kwargs))

def _parse_response(response, dim_names, metric_names):
    rows = []
    for row in response.rows:
        r = {}
        for i, d in enumerate(dim_names): r[d] = row.dimension_values[i].value
        for i, m in enumerate(metric_names):
            v = row.metric_values[i].value
            r[m] = float(v) if '.' in v else int(v)
        rows.append(r)
    return pd.DataFrame(rows)

@st.cache_data(ttl=600, show_spinner=False)
def fetch_traffic(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['date'], ['sessions','totalUsers','activeUsers','newUsers','screenPageViews'])
    return _parse_response(r, ['Date'], ['Sessions','Users','Active Users','New Users','Pageviews'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_channels(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['sessionDefaultChannelGroup','date'], ['sessions','totalUsers','activeUsers','newUsers'])
    return _parse_response(r, ['Channel','Date'], ['Sessions','Users','Active Users','New Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_channel_pages(_ct, prop, start, end):
    """Landing-page breakdown per channel — one row per (channel, landing page, date).
    Uses landingPage so sessions are counted once on the page where the session started."""
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['sessionDefaultChannelGroup','landingPage','date'], ['sessions','totalUsers'])
    return _parse_response(r, ['Channel','Page','Date'], ['Sessions','Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_ai_traffic(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['sessionSource','sessionMedium','date'], ['sessions','totalUsers'])
    df = _parse_response(r, ['Source','Medium','Date'], ['Sessions','Users'])
    return df[df['Source'].str.lower().apply(lambda s: any(ai in s for ai in AI_SOURCES))]

@st.cache_data(ttl=600, show_spinner=False)
def fetch_country(_ct, prop, start, end):
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['country','date'], ['sessions','totalUsers','activeUsers','newUsers'])
    return _parse_response(r, ['Country','Date'], ['Sessions','Users','Active Users','New Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_organic_country(_ct, prop, start, end):
    from google.analytics.data_v1beta.types import FilterExpression, Filter
    creds = _get_creds()
    filt = FilterExpression(filter=Filter(field_name="sessionDefaultChannelGroup", string_filter=Filter.StringFilter(value="Organic Search")))
    r = _ga4_report(creds, prop, start, end, ['country','date'], ['sessions','totalUsers','activeUsers','newUsers'], dim_filter=filt)
    return _parse_response(r, ['Country','Date'], ['Sessions','Users','Active Users','New Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_organic_usa_pages(_ct, prop, start, end):
    from google.analytics.data_v1beta.types import FilterExpression, FilterExpressionList, Filter
    creds = _get_creds()
    filt = FilterExpression(and_group=FilterExpressionList(expressions=[
        FilterExpression(filter=Filter(field_name="country", string_filter=Filter.StringFilter(value="United States"))),
        FilterExpression(filter=Filter(field_name="sessionDefaultChannelGroup", string_filter=Filter.StringFilter(value="Organic Search"))),
    ]))
    r = _ga4_report(creds, prop, start, end, ['pagePath','date'], ['sessions','totalUsers'], dim_filter=filt)
    return _parse_response(r, ['Page','Date'], ['Sessions','Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_organic_users(_ct, prop, start, end):
    from google.analytics.data_v1beta.types import FilterExpression, Filter
    creds = _get_creds()
    filt = FilterExpression(filter=Filter(field_name="sessionDefaultChannelGroup", string_filter=Filter.StringFilter(value="Organic Search")))
    r = _ga4_report(creds, prop, start, end, ['date'], ['sessions','totalUsers','newUsers','activeUsers'], dim_filter=filt)
    return _parse_response(r, ['Date'], ['Sessions','Users','New Users','Active Users'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_traffic_pages(_ct, prop, start, end):
    """Landing-page sessions for overall-traffic drilldown.
    Uses landingPage so each session is counted once on its entry page —
    sum of sessions across landing pages = total sessions for the period."""
    creds = _get_creds()
    r = _ga4_report(creds, prop, start, end, ['landingPage','date'], ['sessions','totalUsers','screenPageViews'])
    return _parse_response(r, ['Page','Date'], ['Sessions','Users','Pageviews'])

@st.cache_data(ttl=600, show_spinner=False)
def fetch_hero_totals(_ct, prop, start, end):
    """Single-week totals matching GA4 UI exactly.
    Uses NO date dimension so metrics are properly deduplicated per the week range.
    - sessions/new_users/active_users/pageviews: no-dim total metrics (matches GA4 Acquisition Overview)
    - users: firstUserDefaultChannelGroup totalUsers sum (matches GA4 channel report Total row)
    - organic_*: filtered to sessionDefaultChannelGroup == 'Organic Search'
    """
    creds = _get_creds()
    out = {'sessions': 0, 'users': 0, 'new_users': 0, 'active_users': 0, 'pageviews': 0,
           'organic_sessions': 0, 'organic_users': 0, 'organic_new_users': 0, 'organic_active_users': 0}
    # Total (no dim) — gives the exact GA4 Total row for the period
    try:
        r1 = _ga4_report(creds, prop, start, end, [], ['sessions', 'totalUsers', 'newUsers', 'activeUsers', 'screenPageViews'])
        if r1.rows:
            mv = r1.rows[0].metric_values
            out['sessions'] = int(float(mv[0].value or 0))
            out['users'] = int(float(mv[1].value or 0))
            out['new_users'] = int(float(mv[2].value or 0))
            out['active_users'] = int(float(mv[3].value or 0))
            out['pageviews'] = int(float(mv[4].value or 0))
    except Exception:
        pass
    # Override Total Users with first-user channel-attribution sum (matches GA4 channel report's "Total Users" row)
    try:
        r2 = _ga4_report(creds, prop, start, end, ['firstUserDefaultChannelGroup'], ['totalUsers'])
        df2 = _parse_response(r2, ['Channel'], ['Users'])
        if len(df2):
            out['users'] = int(df2['Users'].sum())
            org2 = df2[df2['Channel'] == 'Organic Search']
            out['organic_users'] = int(org2['Users'].sum()) if len(org2) else 0
    except Exception:
        pass
    # Organic sessions/active/new — filtered no-dim query
    try:
        from google.analytics.data_v1beta.types import FilterExpression, Filter
        filt = FilterExpression(filter=Filter(field_name="sessionDefaultChannelGroup", string_filter=Filter.StringFilter(value="Organic Search")))
        r3 = _ga4_report(creds, prop, start, end, [], ['sessions', 'newUsers', 'activeUsers'], dim_filter=filt)
        if r3.rows:
            mv3 = r3.rows[0].metric_values
            out['organic_sessions'] = int(float(mv3[0].value or 0))
            out['organic_new_users'] = int(float(mv3[1].value or 0))
            out['organic_active_users'] = int(float(mv3[2].value or 0))
    except Exception:
        pass
    return out

# GSC
@st.cache_data(ttl=600, show_spinner=False)
def fetch_gsc(_ct, site_url, start, end, dim='page'):
    from googleapiclient.discovery import build
    creds = _get_creds()
    svc = build('searchconsole', 'v1', credentials=creds)
    resp = svc.searchanalytics().query(siteUrl=site_url, body={'startDate': start, 'endDate': end, 'dimensions': [dim,'date'], 'rowLimit': 10000}).execute()
    rows = []
    for row in resp.get('rows', []):
        rows.append({dim.capitalize(): row['keys'][0], 'Date': row['keys'][1], 'Clicks': row['clicks'], 'Impressions': row['impressions'], 'CTR': row['ctr']*100, 'Position': row['position']})
    return pd.DataFrame(rows)

def _get_creds():
    # On cloud, load OAuth token from secrets
    if _is_cloud():
        try:
            import base64
            if hasattr(st, 'secrets') and 'oauth' in st.secrets and 'token_b64' in st.secrets['oauth']:
                token_bytes = base64.b64decode(st.secrets['oauth']['token_b64'])
                creds = pickle.loads(token_bytes)
                if creds and creds.expired and creds.refresh_token:
                    from google.auth.transport.requests import Request
                    creds.refresh(Request())
                if creds and creds.valid:
                    return creds
                st.error("Cloud OAuth error: token is invalid or could not be refreshed. Falling back to service account.")
        except Exception as e:
            st.error(f"Cloud OAuth error: {e}")
        creds = get_sa_credentials()
        if creds is None:
            st.error("No valid credentials found. Please update the OAuth token in Streamlit secrets.")
            st.stop()
        return creds
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as f: creds = pickle.load(f)
        if creds and creds.expired and creds.refresh_token:
            from google.auth.transport.requests import Request
            creds.refresh(Request())
            with open(TOKEN_FILE, 'wb') as f: pickle.dump(creds, f)
        return creds
    return get_sa_credentials()

# =====================================================
# DEMO DATA
# =====================================================
def gen_demo(weeks):
    np.random.seed(42)
    all_dates = []
    for m, s, _ in weeks:
        all_dates.extend(pd.date_range(m, s))

    # Traffic
    traffic_rows = []
    for d in all_dates:
        traffic_rows.append({'Date': d.strftime('%Y%m%d'), 'Sessions': np.random.randint(300, 500),
            'Users': np.random.randint(200, 380), 'Active Users': np.random.randint(180, 350),
            'New Users': np.random.randint(100, 250), 'Pageviews': np.random.randint(400, 700)})

    # Channels
    channels = ['Organic Search','Direct','Referral','Organic Social','Email','Paid Search']
    ch_rows = []
    for d in all_dates:
        for ch in channels:
            base = {'Organic Search': 130, 'Direct': 50, 'Referral': 20}.get(ch, 10)
            ch_rows.append({'Channel': ch, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(max(1,base-20), base+30),
                'Users': np.random.randint(max(1,base-30), base+10),
                'Active Users': np.random.randint(max(1,base-35), base+5),
                'New Users': np.random.randint(max(1,base//3), base//2+10)})

    # Channel pages (for drilldown)
    pages = ['/','/blog/','/blog/competitive-intelligence-tools/','/platform/','/news-api/',
             '/blog/market-intelligence-vs-market-research/','/pricing/','/blog/swot-analysis/',
             '/case-studies/','/blog/competitor-analysis/','/resources/','/about/','/demo/']
    cp_rows = []
    for d in all_dates:
        for ch in channels:
            for p in pages[:5]:
                cp_rows.append({'Channel': ch, 'Page': p, 'Date': d.strftime('%Y%m%d'),
                    'Sessions': np.random.randint(1, 15), 'Users': np.random.randint(1, 12)})

    # AI Traffic
    ai_sources = ['chatgpt.com','perplexity.ai','gemini.google.com','claude.ai']
    ai_rows = []
    for d in all_dates:
        for src in ai_sources:
            ai_rows.append({'Source': src, 'Medium': 'referral', 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(1, 15), 'Users': np.random.randint(1, 12)})

    # Country
    countries = ['United States','United Kingdom','Germany','France','India','Canada','Netherlands','Australia','Indonesia','Turkey','Sweden','Ireland','Italy','Austria','Finland','Japan','Brazil']
    co_rows = []
    for d in all_dates:
        for c in countries:
            base = {'United States': 80, 'United Kingdom': 30, 'India': 25, 'Germany': 15}.get(c, 8)
            co_rows.append({'Country': c, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(max(1,base-5), base+10),
                'Users': np.random.randint(max(1,base-8), base+5),
                'Active Users': np.random.randint(max(1,base-10), base+3),
                'New Users': np.random.randint(max(1,base//3), base//2+3)})

    # Organic country
    oc_rows = []
    for d in all_dates:
        for c in countries:
            base = {'United States': 40, 'United Kingdom': 15, 'India': 12, 'Germany': 8}.get(c, 4)
            oc_rows.append({'Country': c, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(max(1,base-3), base+5),
                'Users': np.random.randint(max(1,base-4), base+3),
                'Active Users': np.random.randint(max(1,base-5), base+2),
                'New Users': np.random.randint(max(1,base//3), base//2+2)})

    # Organic USA pages
    op_rows = []
    for d in all_dates:
        for p in pages:
            op_rows.append({'Page': p, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(2, 25), 'Users': np.random.randint(1, 20)})

    # Organic users
    ou_rows = []
    for d in all_dates:
        ou_rows.append({'Date': d.strftime('%Y%m%d'), 'Sessions': np.random.randint(100, 180),
            'Users': np.random.randint(80, 150), 'New Users': np.random.randint(50, 110),
            'Active Users': np.random.randint(70, 140)})

    # Traffic pages (for drilldown)
    tp_rows = []
    for d in all_dates:
        for p in pages:
            tp_rows.append({'Page': p, 'Date': d.strftime('%Y%m%d'),
                'Sessions': np.random.randint(5, 50), 'Users': np.random.randint(3, 40),
                'Pageviews': np.random.randint(8, 60)})

    return (pd.DataFrame(traffic_rows), pd.DataFrame(ch_rows), pd.DataFrame(ai_rows),
            pd.DataFrame(co_rows), pd.DataFrame(oc_rows), pd.DataFrame(op_rows),
            pd.DataFrame(ou_rows), pd.DataFrame(cp_rows), pd.DataFrame(tp_rows))

# =====================================================
# UI HELPERS
# =====================================================
def fmt(n):
    if pd.isna(n) or n == 0: return "0"
    if abs(n) >= 1_000_000: return f"{n/1e6:.1f}M"
    if abs(n) >= 1_000: return f"{n:,.0f}"
    # Treat all GA4 counts (sessions/users) as integers — even when stored as float after fillna/merge.
    return f"{int(round(n))}"
    return str(int(n))

def pct_change(cur, prev):
    if prev == 0: return 0
    return ((cur - prev) / prev) * 100

def change_html(val, fmt_str="+.1f", suffix="%", invert=False):
    positive = val >= 0 if not invert else val <= 0
    cls = "pos" if positive else "neg"
    arrow = "▲" if val > 0 else ("▼" if val < 0 else "–")
    return f'<span class="{cls}">{arrow} {abs(val):{fmt_str}}{suffix}</span>'

def _strip_html_inline(text):
    import re
    return re.sub(r'<[^>]+>', '', str(text)).strip()

def insight_box(items, section_key=None):
    """Render insight box with optional edit button. If section_key is given, insights are editable."""
    if not items: return
    clean_items = [i for i in items if i]
    if not clean_items: return

    # If no key, just render static
    if not section_key:
        bullets = "".join(f"<li>{i}</li>" for i in clean_items)
        st.markdown(f'<div class="insight-box"><strong>💡 Key Insights</strong><ul>{bullets}</ul></div>', unsafe_allow_html=True)
        return

    # Editable insights
    edit_key = f'{section_key}_editing'
    text_key = f'{section_key}_insights_text'
    area_key = f'{section_key}_edit_area'
    last_auto_key = f'{section_key}_last_auto_text'
    user_edited_key = f'{section_key}_user_edited'

    auto_text = "\n".join(_strip_html_inline(i) for i in clean_items)

    if edit_key not in st.session_state:
        st.session_state[edit_key] = False
    if user_edited_key not in st.session_state:
        st.session_state[user_edited_key] = False
    if text_key not in st.session_state:
        st.session_state[text_key] = auto_text
        st.session_state[last_auto_key] = auto_text
    else:
        # Auto-refresh when underlying data changed AND user hasn't manually edited
        if not st.session_state[user_edited_key] and st.session_state.get(last_auto_key) != auto_text:
            st.session_state[text_key] = auto_text
            st.session_state[last_auto_key] = auto_text

    def _toggle():
        st.session_state[edit_key] = not st.session_state[edit_key]
    def _save():
        st.session_state[text_key] = st.session_state[area_key]
        st.session_state[user_edited_key] = True
        st.session_state[edit_key] = False
    def _reset():
        st.session_state[text_key] = auto_text
        st.session_state[last_auto_key] = auto_text
        st.session_state[user_edited_key] = False
        st.session_state[edit_key] = False

    if st.session_state[edit_key]:
        st.text_area("Edit insights (one per line):", value=st.session_state[text_key], height=120, key=area_key)
        c1, c2, c3 = st.columns(3)
        with c1: st.button("Save", use_container_width=True, key=f"{section_key}_save", on_click=_save)
        with c2: st.button("Reset to Auto", use_container_width=True, key=f"{section_key}_reset", on_click=_reset)
        with c3: st.button("Cancel", use_container_width=True, key=f"{section_key}_cancel", on_click=_toggle)
    else:
        lines = st.session_state[text_key].strip().split("\n")
        display_items = [l.strip().lstrip("- ") for l in lines if l.strip()]
        bullets = "".join(f"<li>{i}</li>" for i in display_items)
        st.markdown(f'<div class="insight-box"><strong>💡 Key Insights</strong><ul>{bullets}</ul></div>', unsafe_allow_html=True)
        st.button("Edit Insights", key=f"{section_key}_edit_btn", on_click=_toggle)

def section_start(title, icon="📊"):
    st.markdown(f'<div class="section-card"><div class="section-title">{icon} {title}</div>', unsafe_allow_html=True)

def section_end():
    st.markdown('</div>', unsafe_allow_html=True)

def make_chart(df, x, y, title, chart_type='line', color=None, height=380, show_labels=True):
    colors = ['#1a56db','#0f3460','#10b981','#f59e0b','#ef4444','#8b5cf6','#ec4899']
    if chart_type == 'line':
        fig = px.line(df, x=x, y=y, title=title, markers=True, color=color, color_discrete_sequence=colors,
                      text=y if (show_labels and color is None) else None)
        fig.update_traces(line=dict(width=2.5))
        if show_labels and color is None:
            fig.update_traces(textposition='top center', textfont=dict(size=11, color='#0f3460', family='Inter'),
                              texttemplate='%{text:,.0f}')
        elif show_labels and color is not None:
            fig.update_traces(textposition='top center', textfont=dict(size=10), texttemplate='%{y:,.0f}')
    elif chart_type == 'bar':
        fig = px.bar(df, x=x, y=y, title=title, color=color, color_discrete_sequence=colors, barmode='group')
    # Clean hover labels — show region/category name for multi-line, just value for single
    if color:
        for trace in fig.data:
            trace.hovertemplate = f'<b>{trace.name}</b>: %{{y:,.0f}}<extra></extra>'
    else:
        fig.update_traces(hovertemplate='%{y:,.0f}<extra></extra>')
    fig.update_layout(
        height=height, plot_bgcolor='white', paper_bgcolor='white',
        font=dict(family='Inter', color='#1a1a2e', size=12),
        margin=dict(t=80 if color else 50, b=50, l=60, r=30),
        title=dict(x=0, xanchor='left', y=0.97, yanchor='top', font=dict(size=14)),
        xaxis=dict(gridcolor='#f0f0f0', tickfont=dict(size=11), title=None),
        yaxis=dict(gridcolor='#f0f0f0', tickfont=dict(size=11), rangemode='tozero', title=None,
                   tickformat=',d', separatethousands=True),
        legend=dict(orientation='h', yanchor='top', y=-0.15, x=0, xanchor='left', bgcolor='rgba(0,0,0,0)'),
        hovermode='x unified',
    )
    return fig

def drilldown_table(df, title="Top Pages by Sessions", total_sessions=None):
    """Render a page-level drilldown table.

    GA4 `sessions` broken down by pagePath counts a session once per page visited,
    so summing top-N pages overstates the actual session total. When `total_sessions`
    is passed (the true week/channel total from a no-dim or single-dim query),
    % share is computed against it. Otherwise falls back to the sum of top-N.
    """
    if df.empty:
        st.info("No page-level data available")
        return
    df = df[~df['Page'].astype(str).str.strip().str.lower().isin(['(not set)', '(not provided)', ''])]
    top = df.groupby('Page').agg({'Sessions':'sum','Users':'sum'}).reset_index().sort_values('Sessions', ascending=False).head(10)
    denom = total_sessions if (total_sessions and total_sessions > 0) else top['Sessions'].sum()
    top['% Share'] = (top['Sessions'] / denom * 100).round(1) if denom > 0 else 0
    rows_html = ""
    for i, (_, r) in enumerate(top.iterrows(), 1):
        rows_html += f"<tr><td>{i}</td><td><strong>{r['Page']}</strong></td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Users'])}</td><td>{r['% Share']:.1f}%</td></tr>"
    st.markdown(f'<table class="drilldown-table"><tr><th>#</th><th>Page</th><th>Sessions</th><th>Users</th><th>% Share</th></tr>{rows_html}</table>', unsafe_allow_html=True)

# =====================================================
# SIDEBAR
# =====================================================
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    config = load_config()

    if _is_cloud():
        auth_method = "service_account"
        st.success("☁️ Streamlit Cloud — Service Account")
    else:
        auth_method = st.radio("Data Source", ["demo", "oauth", "service_account"],
            format_func=lambda x: {"demo": "🎯 Demo Data", "oauth": "🔑 Google OAuth", "service_account": "📄 Service Account"}[x],
            index=["demo","oauth","service_account"].index(config.get('auth_method','demo')))
        config['auth_method'] = auth_method; save_config(config)

        if auth_method == "oauth":
            if os.path.exists(TOKEN_FILE):
                st.success("Signed in")
                if st.button("Sign Out & Switch"): os.remove(TOKEN_FILE); st.cache_data.clear(); st.rerun()
            else:
                if st.button("🔑 Sign In with Google", type="primary", use_container_width=True):
                    creds = get_oauth_credentials()
                    if creds: st.rerun()

    st.markdown("---")
    ga4_property = st.text_input("GA4 Property ID", value=config.get('ga4_property','250811000'))
    gsc_site = st.text_input("GSC Site URL", value=config.get('gsc_site',''))
    config['ga4_property'] = ga4_property; config['gsc_site'] = gsc_site; save_config(config)

    st.markdown("---")
    st.markdown("**📅 Date Range**")
    today = datetime.now().date()
    last_sun = today - timedelta(days=(today.weekday() + 1) % 7)
    if last_sun == today: last_sun -= timedelta(days=7)

    date_preset = st.selectbox("Period", [
        "Last week",
        "Last 2 weeks",
        "Last 4 weeks",
        "Last 5 weeks (default)",
        "Last 8 weeks",
        "Last 12 weeks",
        "Custom range",
    ], index=3)

    preset_map = {"Last week": 1, "Last 2 weeks": 2, "Last 4 weeks": 4,
                  "Last 5 weeks (default)": 5, "Last 8 weeks": 8, "Last 12 weeks": 12}

    if date_preset == "Custom range":
        custom_end = st.date_input("End Date", value=today)
        num_weeks = st.number_input("Number of weeks", value=5, min_value=1, max_value=24, step=1)
    else:
        custom_end = None
        num_weeks = preset_map[date_preset]

    # Show selected date range preview
    _preview_weeks = get_last_n_weeks(num_weeks, ref_date=custom_end)
    st.caption(f"📅 {_preview_weeks[0][0].strftime('%d %b %Y')} → {_preview_weeks[-1][1].strftime('%d %b %Y')}")

    st.markdown("---")
    st.markdown("**🎯 Weekly Targets**")
    targets = config.get('targets', DEFAULT_CONFIG['targets'])
    targets['overall_traffic'] = st.number_input("Overall Traffic Target", value=targets.get('overall_traffic', 5538), step=100)
    targets['organic_traffic'] = st.number_input("Organic Traffic Target", value=targets.get('organic_traffic', 1154), step=100)
    targets['overall_users'] = st.number_input("Overall Total Users Target", value=targets.get('overall_users', 4562), step=100)
    targets['organic_users'] = st.number_input("Organic Total Users Target", value=targets.get('organic_users', 706), step=50)
    targets['mqls'] = st.number_input("MQLs Target", value=targets.get('mqls', 13), step=1)
    mqls_achieved = st.number_input("MQLs Achieved (This Week)", value=0, step=1)
    mqls_last_week = st.number_input("MQLs Achieved (Last Week)", value=0, step=1)
    config['targets'] = targets; save_config(config)

    st.markdown("---")
    if st.button("🔄 Refresh Data", use_container_width=True, type="primary"):
        st.cache_data.clear(); st.rerun()

# =====================================================
# WEEKS — current week first in display
# =====================================================
weeks = get_last_n_weeks(num_weeks, ref_date=custom_end)
current_week = weeks[-1]
prev_week = weeks[-2]
full_start = weeks[0][0].strftime('%Y-%m-%d')
full_end = weeks[-1][1].strftime('%Y-%m-%d')
cur_idx = len(weeks) - 1
prev_idx = len(weeks) - 2

# =====================================================
# HEADER
# =====================================================
tag = "DEMO" if auth_method == "demo" else "LIVE"
tag_cls = "tag-demo" if auth_method == "demo" else "tag-live"
st.markdown(f"""
<div class="dashboard-header">
    <h1>Weekly SEO Review <span class="{tag_cls}">{tag}</span></h1>
    <p>📅 Current Week: {current_week[2]} &nbsp;|&nbsp; Range: {weeks[0][0].strftime('%d %b')} – {weeks[-1][1].strftime('%d %b %Y')} ({num_weeks}w) &nbsp;|&nbsp; contify.com</p>
</div>
""", unsafe_allow_html=True)

# =====================================================
# EXPORT HELPERS — PDF & Word
# =====================================================
def _collect_report_data():
    """Gather all report data into a dict for export after data is loaded."""
    return st.session_state.get('_report_data', {})

def _pdf_safe(s):
    """Replace common unicode chars and strip anything outside latin-1 for fpdf Helvetica."""
    if s is None: return ""
    s = str(s)
    repl = {'\u2014':'-', '\u2013':'-', '\u2212':'-', '\u2010':'-', '\u2011':'-',
            '\u2018':"'", '\u2019':"'", '\u201c':'"', '\u201d':'"', '\u2026':'...',
            '\u2022':'*', '\u2192':'->', '\u2190':'<-', '\u25b2':'^', '\u25bc':'v',
            '\u2713':'v', '\u2717':'x', '\u00a0':' ', '\u00b0':'deg'}
    for k, v in repl.items():
        s = s.replace(k, v)
    return s.encode('latin-1', errors='replace').decode('latin-1')

def generate_pdf(report, charts=None):
    """Generate a compact dashboard-style PDF with KPI cards, action box, charts and section tables."""
    from fpdf import FPDF
    import io as _ioc
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    PAGE_W = 190
    NAVY = (15, 52, 96)
    BLUE = (26, 86, 219)
    LIGHT_BLUE_BG = (239, 246, 255)
    LIGHT_GRAY_BG = (248, 250, 252)
    BORDER = (219, 234, 254)
    DARK_TEXT = (30, 41, 59)
    MUTED = (107, 114, 128)
    BADGE_GREEN = (220, 252, 231); GREEN_TEXT = (22, 101, 52)
    BADGE_YELLOW = (254, 243, 199); YELLOW_TEXT = (146, 64, 14)
    BADGE_RED = (254, 226, 226); RED_TEXT = (153, 27, 27)

    def _section_title(title):
        pdf.set_font('Helvetica', 'B', 12)
        pdf.set_text_color(*NAVY)
        pdf.set_x(10)
        pdf.cell(0, 6, _pdf_safe(title), ln=True)
        pdf.set_draw_color(*BLUE)
        pdf.set_line_width(0.6)
        y = pdf.get_y() + 0.5
        pdf.line(10, y, 32, y)
        pdf.set_line_width(0.2)
        pdf.set_draw_color(229, 231, 235)
        pdf.ln(2)
        pdf.set_text_color(*DARK_TEXT)

    # ===== HEADER (compact) =====
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, 210, 24, 'F')
    pdf.set_fill_color(*BLUE)
    pdf.rect(0, 22, 210, 2, 'F')
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 15)
    pdf.set_xy(10, 6)
    pdf.cell(0, 6, _pdf_safe('Weekly SEO Review - Contify'), ln=True)
    pdf.set_font('Helvetica', '', 8)
    pdf.set_xy(10, 14)
    pdf.cell(0, 4, _pdf_safe(f"Week: {report.get('week_label', '')}     Generated: {datetime.now().strftime('%B %d, %Y')}"), ln=True)
    pdf.set_y(28)
    pdf.set_text_color(*DARK_TEXT)

    # ===== KPI CARDS =====
    _section_title('KPI Summary')
    kpi_rows = report.get('kpi_rows', [])
    n_cards = max(len(kpi_rows), 1)
    gap = 1.5
    card_w = (PAGE_W - gap * (n_cards - 1)) / n_cards if n_cards > 1 else PAGE_W
    card_h = 21
    start_x = 10
    start_y = pdf.get_y()
    for i, row in enumerate(kpi_rows):
        x = start_x + i * (card_w + gap)
        pdf.set_fill_color(255, 255, 255)
        pdf.set_draw_color(*BORDER)
        pdf.rect(x, start_y, card_w, card_h, 'DF')
        pdf.set_fill_color(*BLUE)
        pdf.rect(x, start_y, card_w, 1, 'F')
        pdf.set_xy(x + 1.5, start_y + 1.8)
        pdf.set_font('Helvetica', 'B', 6.5)
        pdf.set_text_color(*MUTED)
        pdf.cell(card_w - 3, 3, _pdf_safe(str(row[0]).upper()), 0, 0, 'L')
        pdf.set_xy(x + 1.5, start_y + 6)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(card_w - 3, 6, _pdf_safe(str(row[2])), 0, 0, 'L')
        pdf.set_xy(x + 1.5, start_y + 12.5)
        pdf.set_font('Helvetica', '', 6.5)
        pdf.set_text_color(*MUTED)
        pdf.cell(card_w - 3, 3, _pdf_safe(f"Target: {row[1]}"), 0, 0, 'L')
        try: pct_num = float(str(row[3]).replace('%','').strip())
        except Exception: pct_num = 0
        if pct_num >= 100: badge_fill, badge_text = BADGE_GREEN, GREEN_TEXT
        elif pct_num >= 70: badge_fill, badge_text = BADGE_YELLOW, YELLOW_TEXT
        else: badge_fill, badge_text = BADGE_RED, RED_TEXT
        badge_y = start_y + 16.5
        pdf.set_fill_color(*badge_fill)
        pdf.rect(x + 1.5, badge_y, card_w - 3, 3.5, 'F')
        pdf.set_xy(x + 1.5, badge_y)
        pdf.set_font('Helvetica', 'B', 6.5)
        pdf.set_text_color(*badge_text)
        pdf.cell(card_w - 3, 3.5, _pdf_safe(f"{row[3]} of target"), 0, 0, 'C')
    pdf.set_y(start_y + card_h + 4)
    pdf.set_text_color(*DARK_TEXT)

    # ===== STRATEGIC ACTIONS BOX =====
    insights_list = [ln for ln in report.get('kpi_insights', []) if str(ln).strip()]
    if insights_list:
        _section_title('Strategic Actions & Insights')
        box_x = 10
        box_y = pdf.get_y()
        line_h = 4.2
        box_h = len(insights_list) * line_h + 3
        pdf.set_fill_color(*LIGHT_BLUE_BG)
        pdf.rect(box_x, box_y, PAGE_W, box_h, 'F')
        pdf.set_fill_color(*BLUE)
        pdf.rect(box_x, box_y, 1.5, box_h, 'F')
        pdf.set_xy(box_x + 4, box_y + 1.5)
        pdf.set_font('Helvetica', '', 8.5)
        pdf.set_text_color(*DARK_TEXT)
        for line in insights_list:
            txt = _pdf_safe(f"- {line}")
            pdf.set_x(box_x + 4)
            try: pdf.multi_cell(PAGE_W - 6, line_h, txt)
            except Exception: pdf.cell(PAGE_W - 6, line_h, txt[:160], 0, 1, 'L')
        pdf.set_y(box_y + box_h + 3)

    # ===== CHARTS =====
    if charts:
        pdf.ln(1)
        _section_title('Charts & Visualizations')
        for ctitle, png_bytes in charts:
            try:
                if pdf.get_y() > 215:
                    pdf.add_page()
                pdf.set_fill_color(*LIGHT_BLUE_BG)
                pdf.rect(10, pdf.get_y(), PAGE_W, 5.5, 'F')
                pdf.set_fill_color(*BLUE)
                pdf.rect(10, pdf.get_y(), 1.5, 5.5, 'F')
                pdf.set_xy(13, pdf.get_y() + 0.8)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(*NAVY)
                pdf.cell(0, 4, _pdf_safe(ctitle), ln=True)
                pdf.ln(0.5)
                pdf.set_text_color(*DARK_TEXT)
                pdf.image(_ioc.BytesIO(png_bytes), x=10, w=190)
                pdf.ln(2)
            except Exception:
                pdf.set_font('Helvetica', 'I', 7.5)
                pdf.set_text_color(*MUTED)
                pdf.cell(0, 4, _pdf_safe(f"[chart unavailable: {ctitle}]"), ln=True)
                pdf.set_text_color(*DARK_TEXT)

    # ===== SECTION TABLES =====
    MAX_COLS = 9
    for section in report.get('sections', []):
        if pdf.get_y() > 250:
            pdf.add_page()
        else:
            pdf.ln(2)
        _section_title(section['title'])

        if section.get('table'):
            cols = list(section['table']['columns'])
            data = [list(r) for r in section['table']['data']]
            if len(cols) > MAX_COLS:
                cols = cols[:MAX_COLS]
                data = [r[:MAX_COLS] for r in data]
            n = len(cols)
            if n > 0:
                cw = max(int(PAGE_W / n), 20)
                max_chars = max(int(cw / 1.5), 4)
                pdf.set_font('Helvetica', 'B', 8)
                pdf.set_fill_color(*LIGHT_BLUE_BG)
                pdf.set_text_color(*NAVY)
                pdf.set_draw_color(*BORDER)
                for c in cols:
                    pdf.cell(cw, 8, _pdf_safe(str(c).upper())[:max_chars], 1, 0, 'C', True)
                pdf.ln()
                pdf.set_text_color(*DARK_TEXT)
                pdf.set_font('Helvetica', '', 8)
                for ri, row in enumerate(data):
                    if ri % 2 == 0:
                        pdf.set_fill_color(*LIGHT_GRAY_BG)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    for val in row:
                        pdf.cell(cw, 6.5, _pdf_safe(str(val))[:max_chars], 1, 0, 'C', True)
                    pdf.ln()

        if section.get('insights'):
            pdf.ln(2)
            pdf.set_font('Helvetica', 'I', 8)
            for ins in section['insights']:
                txt = _pdf_safe(f"  * {ins}")
                if not txt.strip(): continue
                pdf.set_x(10)
                try: pdf.multi_cell(PAGE_W, 4, txt)
                except Exception: pdf.cell(PAGE_W, 4, txt[:160], 0, 1, 'L')

        pdf.ln(3)

    out = pdf.output()
    if isinstance(out, (bytearray, memoryview)): out = bytes(out)
    elif isinstance(out, str): out = out.encode('latin-1')
    return out

def generate_docx(report):
    """Generate a Word document report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading('Weekly SEO Review — Contify', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(15, 52, 96)

    doc.add_paragraph(f"Week: {report.get('week_label', '')}  |  Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # KPI Summary
    doc.add_heading('KPI Summary', level=1)
    kpi_rows = report.get('kpi_rows', [])
    if kpi_rows:
        table = doc.add_table(rows=1 + len(kpi_rows), cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        headers = ['Metric', 'Target', 'Achieved', '% Achieved']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        for ri, row in enumerate(kpi_rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    # KPI Insights
    doc.add_paragraph("")
    for line in report.get('kpi_insights', []):
        doc.add_paragraph(line, style='List Bullet')

    # Section tables
    for section in report.get('sections', []):
        doc.add_heading(section['title'], level=1)

        if section.get('table'):
            cols = section['table']['columns']
            data = section['table']['data']
            table = doc.add_table(rows=1 + min(len(data), 15), cols=len(cols))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, c in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = str(c)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            for ri, row in enumerate(data[:15]):
                for ci, val in enumerate(row):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = str(val)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8)

        if section.get('insights'):
            doc.add_paragraph("")
            for ins in section['insights']:
                doc.add_paragraph(ins, style='List Bullet')

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# =====================================================
# FETCH ALL DATA
# =====================================================
with st.spinner("Loading data..."):
    if auth_method == "demo":
        d_traffic, d_channels, d_ai, d_country, d_org_country, d_org_usa_pages, d_org_users, d_ch_pages, d_traffic_pages = gen_demo(weeks)
    else:
        ct = os.path.getmtime(TOKEN_FILE) if os.path.exists(TOKEN_FILE) else 0
        try:
            d_traffic = fetch_traffic(ct, ga4_property, full_start, full_end)
            d_channels = fetch_channels(ct, ga4_property, full_start, full_end)
            d_ai = fetch_ai_traffic(ct, ga4_property, full_start, full_end)
            d_country = fetch_country(ct, ga4_property, full_start, full_end)
            d_org_country = fetch_organic_country(ct, ga4_property, full_start, full_end)
            d_org_usa_pages = fetch_organic_usa_pages(ct, ga4_property, full_start, full_end)
            d_org_users = fetch_organic_users(ct, ga4_property, full_start, full_end)
            d_ch_pages = fetch_channel_pages(ct, ga4_property, full_start, full_end)
            d_traffic_pages = fetch_traffic_pages(ct, ga4_property, full_start, full_end)
        except Exception as e:
            st.error(f"API Error: {e}")
            st.stop()

# Assign weeks
t = assign_weeks(d_traffic, weeks)
ch = assign_weeks(d_channels, weeks)
ai = assign_weeks(d_ai, weeks)
co = assign_weeks(d_country, weeks)
oc = assign_weeks(d_org_country, weeks)
op = assign_weeks(d_org_usa_pages, weeks)
ou = assign_weeks(d_org_users, weeks)
cp = assign_weeks(d_ch_pages, weeks)
tp = assign_weeks(d_traffic_pages, weeks)

# Weekly aggregations
ch_weekly = ch.groupby(['Week_Idx','Week_Label','Channel']).agg({'Sessions':'sum','Users':'sum','Active Users':'sum','New Users':'sum'}).reset_index().sort_values('Week_Idx')

# t_weekly is the SINGLE SOURCE OF TRUTH for total Sessions/Users/etc.
# In live mode each week's row is fetched directly from GA4 with NO date dim, so values match
# GA4 UI Acquisition Overview Total row exactly (= hero cards = chart = sessions table all consistent).
# Demo mode uses the date-level groupby fallback.
if auth_method == 'demo':
    t_weekly = t.groupby(['Week_Idx','Week_Label']).agg({'Sessions':'sum','Users':'sum','Active Users':'sum','New Users':'sum','Pageviews':'sum'}).reset_index().sort_values('Week_Idx')
else:
    _ct_w = os.path.getmtime(TOKEN_FILE) if os.path.exists(TOKEN_FILE) else 0
    _rows = []
    for _w_idx, (_w_start, _w_end, _w_label) in enumerate(weeks):
        _h = fetch_hero_totals(_ct_w, ga4_property, _w_start.strftime('%Y-%m-%d'), _w_end.strftime('%Y-%m-%d'))
        _rows.append({
            'Week_Idx': _w_idx, 'Week_Label': _w_label,
            'Sessions': _h['sessions'], 'Users': _h['users'],
            'Active Users': _h['active_users'], 'New Users': _h['new_users'],
            'Pageviews': _h['pageviews'],
        })
    t_weekly = pd.DataFrame(_rows).sort_values('Week_Idx')
ai_weekly = ai.groupby(['Week_Idx','Week_Label','Source']).agg({'Sessions':'sum','Users':'sum'}).reset_index()
co_weekly = co.groupby(['Week_Idx','Week_Label','Country']).agg({'Sessions':'sum','Users':'sum','Active Users':'sum','New Users':'sum'}).reset_index()
oc_weekly = oc.groupby(['Week_Idx','Week_Label','Country']).agg({'Sessions':'sum','Users':'sum','Active Users':'sum','New Users':'sum'}).reset_index()
op_weekly = op.groupby(['Week_Idx','Week_Label','Page']).agg({'Sessions':'sum','Users':'sum'}).reset_index()
ou_weekly = ou.groupby(['Week_Idx','Week_Label']).agg({'Sessions':'sum','Users':'sum','New Users':'sum','Active Users':'sum'}).reset_index().sort_values('Week_Idx')
cp_weekly = cp.groupby(['Week_Idx','Week_Label','Channel','Page']).agg({'Sessions':'sum','Users':'sum'}).reset_index()
tp_weekly = tp.groupby(['Week_Idx','Week_Label','Page']).agg({'Sessions':'sum','Users':'sum','Pageviews':'sum'}).reset_index()

# =====================================================
# HERO METRIC CARDS — Top-Level Summary
# =====================================================
_cur_t = t_weekly[t_weekly.Week_Idx==cur_idx]
_prev_t = t_weekly[t_weekly.Week_Idx==prev_idx]

# Hero cards read from t_weekly — same dataframe used by the trend chart and sessions table,
# so all three displays show identical numbers (matching GA4 UI Total).
cur_sessions = int(_cur_t['Sessions'].sum()) if len(_cur_t) else 0
cur_users = int(_cur_t['Users'].sum()) if len(_cur_t) else 0
prev_sessions = int(_prev_t['Sessions'].sum()) if len(_prev_t) else 0
prev_users = int(_prev_t['Users'].sum()) if len(_prev_t) else 0

if auth_method == 'demo':
    # Demo fallback uses channel-sum for organic
    organic_cur = ch_weekly[(ch_weekly.Week_Idx==cur_idx) & (ch_weekly.Channel=='Organic Search')]
    organic_prev = ch_weekly[(ch_weekly.Week_Idx==prev_idx) & (ch_weekly.Channel=='Organic Search')]
    org_sessions = int(organic_cur['Sessions'].sum()) if len(organic_cur) else 0
    org_users = int(organic_cur['Users'].sum()) if len(organic_cur) else 0
    org_sessions_prev = int(organic_prev['Sessions'].sum()) if len(organic_prev) else 0
    org_users_prev = int(organic_prev['Users'].sum()) if len(organic_prev) else 0
else:
    # Organic hero — use the same fetch_hero_totals (cached) so organic also matches GA4 UI exactly.
    _ct_h = os.path.getmtime(TOKEN_FILE) if os.path.exists(TOKEN_FILE) else 0
    _h_cur = fetch_hero_totals(_ct_h, ga4_property, current_week[0].strftime('%Y-%m-%d'), current_week[1].strftime('%Y-%m-%d'))
    _h_prev = fetch_hero_totals(_ct_h, ga4_property, prev_week[0].strftime('%Y-%m-%d'), prev_week[1].strftime('%Y-%m-%d'))
    org_sessions = _h_cur['organic_sessions']
    org_users = _h_cur['organic_users']
    org_sessions_prev = _h_prev['organic_sessions']
    org_users_prev = _h_prev['organic_users']

def _hero_card(label, value, prev_val):
    # Handle prev=0 edge cases so MQLs (or any new metric) doesn't show misleading "0.0% vs last week"
    if prev_val == 0 and value == 0:
        change_html_str = '<span style="color:#9ca3af;">no prior data</span>'
    elif prev_val == 0 and value > 0:
        change_html_str = '<span style="color:#059669;">▲ new this week</span>'
    else:
        chg = pct_change(value, prev_val)
        arrow = "▲" if chg > 0 else ("▼" if chg < 0 else "–")
        arrow_color = "#059669" if chg >= 0 else "#dc2626"
        change_html_str = f'<span style="color:{arrow_color};">{arrow} {abs(chg):.1f}% vs last week</span>'
    return f'''<div style="background:#ffffff;border:1px solid #e2e8f0;border-radius:14px;padding:20px 16px;text-align:center;
        box-shadow:0 3px 12px rgba(15,52,96,0.06);flex:1;min-width:0;">
        <div style="font-size:0.72rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.6px;font-weight:600;margin-bottom:6px;">{label}</div>
        <div style="font-size:1.8rem;font-weight:800;color:#0f3460;">{fmt(value)}</div>
        <div style="font-size:0.82rem;font-weight:600;margin-top:4px;">{change_html_str}</div>
    </div>'''

st.markdown(f'''<div style="display:flex;gap:16px;margin-bottom:24px;">
    {_hero_card("Total Sessions", cur_sessions, prev_sessions)}
    {_hero_card("Total Users", cur_users, prev_users)}
    {_hero_card("MQLs", mqls_achieved, mqls_last_week)}
    {_hero_card("Organic Sessions", org_sessions, org_sessions_prev)}
    {_hero_card("Organic Users", org_users, org_users_prev)}
</div>''', unsafe_allow_html=True)

# =====================================================
# SECTION 1: KPI SUMMARY
# =====================================================
section_start("KPI Summary — This Week", "🎯")

cur_traffic = _cur_t

kpi_data = [
    ("Overall Traffic", targets['overall_traffic'], cur_sessions),
    ("Organic Traffic", targets['organic_traffic'], org_sessions),
    ("Overall Total Users", targets['overall_users'], cur_users),
    ("Organic Total Users", targets['organic_users'], org_users),
    ("MQLs", targets['mqls'], mqls_achieved),
]

rows_html = ""
for name, target, achieved in kpi_data:
    pct = (achieved / target * 100) if target > 0 else 0
    color = "#059669" if pct >= 100 else ("#f59e0b" if pct >= 80 else "#dc2626")
    bar_w = min(pct, 100)
    status = "Exceeded" if pct >= 100 else ("On Track" if pct >= 80 else "Behind")
    status_bg = "#ecfdf5" if pct >= 100 else ("#fffbeb" if pct >= 80 else "#fef2f2")
    rows_html += f"""<tr>
        <td><strong>{name}</strong></td>
        <td style="text-align:center;">{fmt(target)}</td>
        <td style="text-align:center;"><strong>{fmt(achieved)}</strong></td>
        <td style="text-align:center;"><span style="font-weight:700;color:{color};font-size:1rem;">{pct:.0f}%</span></td>
    </tr>"""

st.markdown(f'<table class="kpi-table"><tr><th>Metric</th><th style="text-align:center;">Target</th><th style="text-align:center;">Achieved</th><th style="text-align:center;">Progress</th></tr>{rows_html}</table>', unsafe_allow_html=True)

# KPI insights — positive tone, editable
kpi_auto_lines = []
for name, target, achieved in kpi_data:
    pct = (achieved / target * 100) if target > 0 else 0
    if pct >= 100:
        kpi_auto_lines.append(f"{name}: Exceeded target — achieved {pct:.0f}% ({fmt(achieved)} / {fmt(target)})")
    elif pct >= 80:
        kpi_auto_lines.append(f"{name}: On track — achieved {pct:.0f}% of target ({fmt(achieved)} / {fmt(target)})")
    else:
        kpi_auto_lines.append(f"{name}: Achieved {pct:.0f}% of target so far ({fmt(achieved)} / {fmt(target)})")

insight_box(kpi_auto_lines, section_key="kpi")

section_end()

# =====================================================
# SECTION 2: OVERALL WEBSITE TRAFFIC — current week first
# =====================================================
section_start("Overall Website Traffic — Last 5 Weeks", "📈")

# Chart — line chart, not area
fig = make_chart(t_weekly, 'Week_Label', 'Sessions', 'Weekly Sessions', chart_type='line')
st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

# Table with WoW change — current week FIRST (reverse order), no Pageviews
display_t = t_weekly[['Week_Idx','Week_Label','Sessions','Users']].copy()
display_t['Sessions Change'] = display_t['Sessions'].pct_change() * 100
display_t = display_t.sort_values('Week_Idx', ascending=False)  # current week first

# Target achievement for each week
rows_html = ""
for _, r in display_t.iterrows():
    chg = change_html(r['Sessions Change']) if pd.notna(r['Sessions Change']) else "–"
    target_pct = (r['Sessions'] / targets['overall_traffic'] * 100) if targets['overall_traffic'] > 0 else 0
    pct_color = "#059669" if target_pct >= 100 else ("#f59e0b" if target_pct >= 80 else "#dc2626")
    current_tag = ' <span class="week-current">CURRENT</span>' if r['Week_Idx'] == cur_idx else ''
    rows_html += f"<tr><td class='week-label-cell'><strong>{r['Week_Label']}</strong>{current_tag}</td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Users'])}</td><td><span style='color:{pct_color};font-weight:600'>{target_pct:.0f}%</span></td><td>{chg}</td></tr>"
st.markdown(f'<table class="change-table"><tr><th>Week</th><th>Sessions</th><th>Users</th><th>vs Target</th><th>WoW Change</th></tr>{rows_html}</table>', unsafe_allow_html=True)

cur_s = t_weekly[t_weekly.Week_Idx==cur_idx]['Sessions'].sum()
prev_s = t_weekly[t_weekly.Week_Idx==prev_idx]['Sessions'].sum()

# Drilldown — toggle to see page source
if st.checkbox("Show page-level breakdown (Current Week)", value=False, key="drill_traffic"):
    cur_pages = tp_weekly[tp_weekly.Week_Idx==cur_idx]
    drilldown_table(cur_pages, "Top Pages — Current Week", total_sessions=cur_s)

chg = pct_change(cur_s, prev_s)
target_ach = (cur_s / targets['overall_traffic'] * 100) if targets['overall_traffic'] > 0 else 0
insight_box([
    f'This week we achieved <strong>{target_ach:.0f}%</strong> of traffic target ({fmt(cur_s)} / {fmt(targets["overall_traffic"])})',
    f'Overall traffic {change_html(chg)} vs last week ({fmt(cur_s)} vs {fmt(prev_s)})',
], section_key="traffic")
section_end()

# =====================================================
# SECTION 3: TRAFFIC BY CHANNELS
# =====================================================
section_start("Traffic by Channels — vs Previous Week", "📊")

cur_ch = ch_weekly[ch_weekly.Week_Idx==cur_idx].groupby('Channel')['Sessions'].sum().reset_index()
prev_ch = ch_weekly[ch_weekly.Week_Idx==prev_idx].groupby('Channel')['Sessions'].sum().reset_index()
merged_ch = cur_ch.merge(prev_ch, on='Channel', suffixes=(' This Week',' Last Week'), how='outer').fillna(0)
merged_ch['Change'] = merged_ch['Sessions This Week'] - merged_ch['Sessions Last Week']
merged_ch['% Change'] = merged_ch.apply(lambda r: pct_change(r['Sessions This Week'], r['Sessions Last Week']), axis=1)
merged_ch = merged_ch.sort_values('Sessions This Week', ascending=False)

rows_html = ""
for _, r in merged_ch.iterrows():
    rows_html += f"<tr><td><strong>{r['Channel']}</strong></td><td>{fmt(r['Sessions This Week'])}</td><td>{fmt(r['Sessions Last Week'])}</td><td>{change_html(r['Change'], '.0f', '')}</td><td>{change_html(r['% Change'])}</td></tr>"
st.markdown(f'<table class="change-table"><tr><th>Channel</th><th>This Week</th><th>Last Week</th><th>Change</th><th>% Change</th></tr>{rows_html}</table>', unsafe_allow_html=True)

# Drilldown per channel
if st.checkbox("Show page breakdown per channel", value=False, key="drill_channel"):
    sel_channel = st.selectbox("Select Channel", merged_ch['Channel'].tolist(), key="ch_drill")
    ch_pages = cp_weekly[(cp_weekly.Week_Idx==cur_idx) & (cp_weekly.Channel==sel_channel)]
    _ch_total = merged_ch[merged_ch['Channel'] == sel_channel]['Sessions This Week'].sum() if len(merged_ch) else 0
    drilldown_table(ch_pages, f"Top Pages — {sel_channel}", total_sessions=_ch_total)

ch_insights = []
top_ch = merged_ch.iloc[0] if len(merged_ch) else None
if top_ch is not None:
    ch_insights.append(f'<strong>{top_ch["Channel"]}</strong> leads with {fmt(top_ch["Sessions This Week"])} sessions ({change_html(top_ch["% Change"])})')
biggest_gain = merged_ch[merged_ch['% Change']>0].nlargest(1, '% Change')
if len(biggest_gain):
    r = biggest_gain.iloc[0]
    ch_insights.append(f'<strong>{r["Channel"]}</strong> had the biggest growth: {change_html(r["% Change"])}')
insight_box(ch_insights, section_key="channels")
section_end()

# =====================================================
# SECTION 4: AI TRAFFIC
# =====================================================
section_start("AI Traffic", "🤖")

ai_cur = ai_weekly[ai_weekly.Week_Idx==cur_idx].groupby('Source').agg({'Sessions':'sum','Users':'sum'}).reset_index().sort_values('Sessions', ascending=False)
ai_prev_total = ai_weekly[ai_weekly.Week_Idx==prev_idx]['Sessions'].sum()
ai_cur_total = ai_cur['Sessions'].sum()

c1, c2 = st.columns([1, 2])
with c1:
    ai_chg = pct_change(ai_cur_total, ai_prev_total)
    st.markdown(f'<div class="metric-mini"><div class="label">Total AI Traffic</div><div class="val">{fmt(ai_cur_total)}</div><div class="change">{change_html(ai_chg)} vs last week</div></div>', unsafe_allow_html=True)
with c2:
    if len(ai_cur):
        ai_cur['% Share'] = (ai_cur['Sessions'] / ai_cur_total * 100).round(1)
        rows_html = ""
        for _, r in ai_cur.iterrows():
            rows_html += f"<tr><td><strong>{r['Source']}</strong></td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Users'])}</td><td>{r['% Share']:.1f}%</td></tr>"
        st.markdown(f'<table class="change-table"><tr><th>Source</th><th>Sessions</th><th>Users</th><th>% Share</th></tr>{rows_html}</table>', unsafe_allow_html=True)

insight_box([f'AI referral traffic: <strong>{fmt(ai_cur_total)}</strong> sessions ({change_html(ai_chg)} WoW)',
             f'Top AI source: <strong>{ai_cur.iloc[0]["Source"] if len(ai_cur) else "N/A"}</strong>' + (f' with {fmt(ai_cur.iloc[0]["Sessions"])} sessions' if len(ai_cur) else '')], section_key="ai")
section_end()

# =====================================================
# SECTION 5: ORGANIC TRAFFIC
# =====================================================
section_start("Organic Traffic — Last 5 Weeks", "🌿")

org_weekly = ch_weekly[ch_weekly.Channel=='Organic Search'].groupby(['Week_Idx','Week_Label']).agg({'Sessions':'sum','Users':'sum','New Users':'sum'}).reset_index().sort_values('Week_Idx')

fig = make_chart(org_weekly, 'Week_Label', 'Sessions', 'Organic Sessions (Weekly)', chart_type='line')
st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

# Table — current week first
display_org = org_weekly[['Week_Idx','Week_Label','Sessions','Users','New Users']].copy()
display_org['WoW Change'] = display_org['Sessions'].pct_change() * 100
display_org_rev = display_org.sort_values('Week_Idx', ascending=False)
rows_html = ""
for _, r in display_org_rev.iterrows():
    chg = change_html(r['WoW Change']) if pd.notna(r['WoW Change']) else "–"
    target_pct = (r['Sessions'] / targets['organic_traffic'] * 100) if targets['organic_traffic'] > 0 else 0
    pct_color = "#059669" if target_pct >= 100 else ("#f59e0b" if target_pct >= 80 else "#dc2626")
    current_tag = ' <span class="week-current">CURRENT</span>' if r['Week_Idx'] == cur_idx else ''
    rows_html += f"<tr><td><strong>{r['Week_Label']}</strong>{current_tag}</td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Users'])}</td><td>{fmt(r['New Users'])}</td><td><span style='color:{pct_color};font-weight:600'>{target_pct:.0f}%</span></td><td>{chg}</td></tr>"
st.markdown(f'<table class="change-table"><tr><th>Week</th><th>Organic Sessions</th><th>Users</th><th>New Users</th><th>vs Target</th><th>WoW Change</th></tr>{rows_html}</table>', unsafe_allow_html=True)

org_cur_s = org_weekly[org_weekly.Week_Idx==cur_idx]['Sessions'].sum()
org_prev_s = org_weekly[org_weekly.Week_Idx==prev_idx]['Sessions'].sum()
org_chg = pct_change(org_cur_s, org_prev_s)
org_target_ach = (org_cur_s / targets['organic_traffic'] * 100) if targets['organic_traffic'] > 0 else 0
insight_box([
    f'This week we achieved <strong>{org_target_ach:.0f}%</strong> of organic traffic target ({fmt(org_cur_s)} / {fmt(targets["organic_traffic"])})',
    f'Organic traffic {change_html(org_chg)} vs last week',
    f'Organic share of total: <strong>{(org_cur_s/cur_sessions*100):.1f}%</strong>' if cur_sessions > 0 else '',
], section_key="organic")
section_end()

# =====================================================
# SECTION 6: TOP ORGANIC COUNTRIES
# =====================================================
section_start("Top Organic Countries — vs Last Week", "🌍")

oc_cur = oc_weekly[oc_weekly.Week_Idx==cur_idx].groupby('Country')['Sessions'].sum().nlargest(5).reset_index()
oc_prev = oc_weekly[oc_weekly.Week_Idx==prev_idx].groupby('Country')['Sessions'].sum().reset_index()
oc_merged = oc_cur.merge(oc_prev, on='Country', suffixes=(' This Week',' Last Week'), how='left').fillna(0)
oc_merged['Change %'] = oc_merged.apply(lambda r: pct_change(r['Sessions This Week'], r['Sessions Last Week']), axis=1)

rows_html = ""
for _, r in oc_merged.iterrows():
    rows_html += f"<tr><td><strong>{r['Country']}</strong></td><td>{fmt(r['Sessions This Week'])}</td><td>{fmt(r['Sessions Last Week'])}</td><td>{change_html(r['Change %'])}</td></tr>"
st.markdown(f'<table class="change-table"><tr><th>Country</th><th>This Week</th><th>Last Week</th><th>Change</th></tr>{rows_html}</table>', unsafe_allow_html=True)

oc_insights = []
for _, r in oc_merged.head(3).iterrows():
    oc_insights.append(f'<strong>{r["Country"]}</strong>: {fmt(r["Sessions This Week"])} organic sessions ({change_html(r["Change %"])})')
insight_box(oc_insights, section_key="org_countries")
section_end()

# =====================================================
# SECTION 7: OVERALL TOTAL USERS
# =====================================================
section_start("Overall Total Users — Last 5 Weeks", "👥")

fig = make_chart(t_weekly, 'Week_Label', 'Users', 'Total Users (Weekly)', chart_type='line')
st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

display_u = t_weekly[['Week_Idx','Week_Label','Users','New Users','Active Users']].copy()
display_u['WoW Change'] = display_u['Users'].pct_change() * 100
display_u_rev = display_u.sort_values('Week_Idx', ascending=False)
rows_html = ""
for _, r in display_u_rev.iterrows():
    chg = change_html(r['WoW Change']) if pd.notna(r['WoW Change']) else "–"
    target_pct = (r['Users'] / targets['overall_users'] * 100) if targets['overall_users'] > 0 else 0
    pct_color = "#059669" if target_pct >= 100 else ("#f59e0b" if target_pct >= 80 else "#dc2626")
    current_tag = ' <span class="week-current">CURRENT</span>' if r['Week_Idx'] == cur_idx else ''
    rows_html += f"<tr><td><strong>{r['Week_Label']}</strong>{current_tag}</td><td>{fmt(r['Users'])}</td><td>{fmt(r['New Users'])}</td><td>{fmt(r['Active Users'])}</td><td><span style='color:{pct_color};font-weight:600'>{target_pct:.0f}%</span></td><td>{chg}</td></tr>"
st.markdown(f'<table class="change-table"><tr><th>Week</th><th>Users</th><th>New Users</th><th>Active Users</th><th>vs Target</th><th>WoW Change</th></tr>{rows_html}</table>', unsafe_allow_html=True)

u_cur = t_weekly[t_weekly.Week_Idx==cur_idx]['Users'].sum()
u_prev = t_weekly[t_weekly.Week_Idx==prev_idx]['Users'].sum()
u_target_ach = (u_cur / targets['overall_users'] * 100) if targets['overall_users'] > 0 else 0
insight_box([
    f'This week we achieved <strong>{u_target_ach:.0f}%</strong> of users target ({fmt(u_cur)} / {fmt(targets["overall_users"])})',
    f'Total users {change_html(pct_change(u_cur, u_prev))} vs last week ({fmt(u_cur)} vs {fmt(u_prev)})',
], section_key="overall_users")
section_end()

# =====================================================
# SECTION 9: TRAFFIC BY REGION
# =====================================================
section_start("Overall Traffic by Region — Last 5 Weeks", "🗺️")

def classify_region(country):
    if country == 'United States': return 'US'
    if country == 'United Kingdom': return 'UK'
    if country in EUROPE_COUNTRIES: return 'Europe'
    if country in EXCLUDE_FROM_OTHER: return '_exclude'
    return 'Other'

co_weekly_copy = co_weekly.copy()
co_weekly_copy['Region'] = co_weekly_copy['Country'].apply(classify_region)
co_weekly_copy = co_weekly_copy[co_weekly_copy.Region != '_exclude']
region_weekly = co_weekly_copy.groupby(['Week_Idx','Week_Label','Region']).agg({'Sessions':'sum'}).reset_index()

fig = make_chart(region_weekly, 'Week_Label', 'Sessions', 'Sessions by Region', chart_type='line', color='Region', show_labels=False)
st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

# Pivot table — current week first
pivot = region_weekly.pivot_table(index=['Week_Idx','Week_Label'], columns='Region', values='Sessions', aggfunc='sum').fillna(0)
for reg in ['US','UK','Europe','Other']:
    if reg not in pivot.columns: pivot[reg] = 0
pivot = pivot[['US','UK','Europe','Other']].reset_index().sort_values('Week_Idx', ascending=False)

def _val_pct(v, total):
    pct = (v / total * 100) if total > 0 else 0
    return f"{fmt(v)} <span style='color:#6b7280;font-size:0.85em;'>({pct:.1f}%)</span>"

# Per-week total across ALL countries (including India/Indonesia/Turkey) so %
# reflects share of total traffic, not share of the visible 4 buckets.
_week_country_total = co_weekly.groupby('Week_Idx')['Sessions'].sum().to_dict()

rows_html = ""
for _, r in pivot.iterrows():
    is_current = r['Week_Idx'] == cur_idx
    current_tag = ' <span class="week-current">CURRENT</span>' if is_current else ''
    if is_current:
        row_total = _week_country_total.get(r['Week_Idx'], 0)
        rows_html += (
            f"<tr><td><strong>{r['Week_Label']}</strong>{current_tag}</td>"
            f"<td>{_val_pct(r['US'], row_total)}</td>"
            f"<td>{_val_pct(r['UK'], row_total)}</td>"
            f"<td>{_val_pct(r['Europe'], row_total)}</td>"
            f"<td>{_val_pct(r['Other'], row_total)}</td></tr>"
        )
    else:
        rows_html += (
            f"<tr><td><strong>{r['Week_Label']}</strong></td>"
            f"<td>{fmt(r['US'])}</td>"
            f"<td>{fmt(r['UK'])}</td>"
            f"<td>{fmt(r['Europe'])}</td>"
            f"<td>{fmt(r['Other'])}</td></tr>"
        )
st.markdown(f'<table class="change-table"><tr><th>Week</th><th>🇺🇸 US</th><th>🇬🇧 UK</th><th>🇪🇺 Europe</th><th>🌍 Other</th></tr>{rows_html}</table>', unsafe_allow_html=True)

# Insights
reg_cur = region_weekly[region_weekly.Week_Idx==cur_idx].set_index('Region')['Sessions']
reg_prev = region_weekly[region_weekly.Week_Idx==prev_idx].set_index('Region')['Sessions']
reg_insights = []
for reg in ['US','UK','Europe','Other']:
    c = reg_cur.get(reg, 0); p = reg_prev.get(reg, 0)
    reg_insights.append(f'<strong>{reg}</strong>: {fmt(c)} sessions ({change_html(pct_change(c, p))})')
insight_box(reg_insights, section_key="region")
section_end()

# =====================================================
# SECTION 10: ORGANIC TRAFFIC USA + TOP PAGES
# =====================================================
section_start("Organic Traffic — USA", "🇺🇸")

us_org = oc_weekly[oc_weekly.Country=='United States'].groupby(['Week_Idx','Week_Label']).agg({'Sessions':'sum'}).reset_index().sort_values('Week_Idx')

fig = make_chart(us_org, 'Week_Label', 'Sessions', 'US Organic Sessions (Weekly)', chart_type='line')
st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

# Top 10 landing pages
st.markdown("**Top 10 Landing Pages — US Organic (This Week)**")
us_pages = op_weekly[op_weekly.Week_Idx==cur_idx].groupby('Page').agg({'Sessions':'sum','Users':'sum'}).reset_index().sort_values('Sessions', ascending=False).head(10)
if len(us_pages):
    us_pages['% Share'] = (us_pages['Sessions'] / us_pages['Sessions'].sum() * 100).round(1)
    rows_html = ""
    for i, (_, r) in enumerate(us_pages.iterrows(), 1):
        rows_html += f"<tr><td>{i}</td><td><strong>{r['Page']}</strong></td><td>{fmt(r['Sessions'])}</td><td>{fmt(r['Users'])}</td><td>{r['% Share']:.1f}%</td></tr>"
    st.markdown(f'<table class="change-table"><tr><th>#</th><th>Page</th><th>Sessions</th><th>Users</th><th>% Share</th></tr>{rows_html}</table>', unsafe_allow_html=True)

us_cur = us_org[us_org.Week_Idx==cur_idx]['Sessions'].sum()
us_prev = us_org[us_org.Week_Idx==prev_idx]['Sessions'].sum()
insight_box([f'US organic traffic: {fmt(us_cur)} sessions ({change_html(pct_change(us_cur, us_prev))})',
             f'Top page: <strong>{us_pages.iloc[0]["Page"]}</strong> with {fmt(us_pages.iloc[0]["Sessions"])} sessions' if len(us_pages) else ''], section_key="us_organic")
section_end()

# =====================================================
# COLLECT REPORT DATA FOR EXPORT
# =====================================================
def _strip_html(text):
    """Remove HTML tags for plain text export."""
    import re
    return re.sub(r'<[^>]+>', '', str(text)).strip()

report_data = {
    'week_label': current_week[2],
    'kpi_rows': [(name, fmt(target), fmt(achieved), f"{(achieved/target*100) if target > 0 else 0:.0f}%") for name, target, achieved in kpi_data],
    'kpi_insights': [line.strip().lstrip("- ") for line in st.session_state.get('kpi_insights_text', '').split("\n") if line.strip()],
    'sections': [],
}

# Section 2: Overall Traffic
traffic_section = {'title': 'Overall Website Traffic', 'table': {'columns': ['Week','Sessions','Users'], 'data': []}, 'insights': []}
for _, r in t_weekly.sort_values('Week_Idx', ascending=False).iterrows():
    traffic_section['table']['data'].append([r['Week_Label'], fmt(r['Sessions']), fmt(r['Users'])])
t_ach = (cur_s / targets['overall_traffic'] * 100) if targets['overall_traffic'] > 0 else 0
_wow_pct = pct_change(cur_s, prev_s)
traffic_section['insights'] = [f"Achieved {t_ach:.0f}% of traffic target ({fmt(cur_s)} / {fmt(targets['overall_traffic'])})", f"WoW change: {_wow_pct:+.1f}%"]
report_data['sections'].append(traffic_section)

# Section 3: Channels
ch_section = {'title': 'Traffic by Channels', 'table': {'columns': ['Channel','This Week','Last Week','Change %'], 'data': []}, 'insights': []}
for _, r in merged_ch.iterrows():
    ch_section['table']['data'].append([r['Channel'], fmt(r['Sessions This Week']), fmt(r['Sessions Last Week']), f"{r['% Change']:+.1f}%"])
if top_ch is not None:
    ch_section['insights'].append(f"{top_ch['Channel']} leads with {fmt(top_ch['Sessions This Week'])} sessions")
report_data['sections'].append(ch_section)

# Section 4: AI Traffic
ai_section = {'title': 'AI Traffic', 'table': {'columns': ['Source','Sessions','Users','% Share'], 'data': []}, 'insights': []}
for _, r in ai_cur.iterrows():
    ai_section['table']['data'].append([r['Source'], fmt(r['Sessions']), fmt(r['Users']), f"{r.get('% Share',0):.1f}%"])
ai_section['insights'] = [f"Total AI traffic: {fmt(ai_cur_total)} sessions"]
report_data['sections'].append(ai_section)

# Section 5: Organic Traffic
org_section = {'title': 'Organic Traffic', 'table': {'columns': ['Week','Sessions','Users','New Users'], 'data': []}, 'insights': []}
for _, r in org_weekly.sort_values('Week_Idx', ascending=False).iterrows():
    org_section['table']['data'].append([r['Week_Label'], fmt(r['Sessions']), fmt(r['Users']), fmt(r['New Users'])])
org_a = (org_cur_s / targets['organic_traffic'] * 100) if targets['organic_traffic'] > 0 else 0
org_section['insights'] = [f"Achieved {org_a:.0f}% of organic target ({fmt(org_cur_s)} / {fmt(targets['organic_traffic'])})"]
report_data['sections'].append(org_section)

# Section 6: Top Organic Countries
oc_section = {'title': 'Top Organic Countries', 'table': {'columns': ['Country','This Week','Last Week','Change'], 'data': []}, 'insights': []}
for _, r in oc_merged.iterrows():
    oc_section['table']['data'].append([r['Country'], fmt(r['Sessions This Week']), fmt(r['Sessions Last Week']), f"{r['Change %']:+.1f}%"])
report_data['sections'].append(oc_section)

# Section 7: Overall Users
user_section = {'title': 'Overall Total Users', 'table': {'columns': ['Week','Users','New Users','Active Users'], 'data': []}, 'insights': []}
for _, r in t_weekly.sort_values('Week_Idx', ascending=False).iterrows():
    user_section['table']['data'].append([r['Week_Label'], fmt(r['Users']), fmt(r['New Users']), fmt(r['Active Users'])])
u_ach = (u_cur / targets['overall_users'] * 100) if targets['overall_users'] > 0 else 0
user_section['insights'] = [f"Achieved {u_ach:.0f}% of users target ({fmt(u_cur)} / {fmt(targets['overall_users'])})"]
report_data['sections'].append(user_section)

# Section 9: Region
reg_section = {'title': 'Traffic by Region', 'table': {'columns': ['Week','US','UK','Europe','Other'], 'data': []}, 'insights': []}
for _, r in pivot.iterrows():
    reg_section['table']['data'].append([r['Week_Label'], fmt(r['US']), fmt(r['UK']), fmt(r['Europe']), fmt(r['Other'])])
report_data['sections'].append(reg_section)

# Section 10: US Organic Pages
us_section = {'title': 'US Organic — Top Landing Pages', 'table': {'columns': ['#','Page','Sessions','Users','% Share'], 'data': []}, 'insights': []}
if len(us_pages):
    for i, (_, r) in enumerate(us_pages.iterrows(), 1):
        us_section['table']['data'].append([i, r['Page'], fmt(r['Sessions']), fmt(r['Users']), f"{r['% Share']:.1f}%"])
report_data['sections'].append(us_section)

# =====================================================
# DOWNLOAD BUTTONS
# =====================================================
st.markdown("---")
st.markdown("### 📥 Download Report")
dl1, dl2, dl3 = st.columns(3)

with dl1:
    try:
        pdf_bytes = generate_pdf(report_data)
        st.download_button(
            label="📄 Download PDF",
            data=pdf_bytes,
            file_name=f"SEO_Weekly_Review_{current_week[0].strftime('%Y%m%d')}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    except ImportError:
        st.warning("Install `fpdf2` for PDF export: `pip install fpdf2`")
    except Exception as e:
        st.error(f"PDF error: {e}")

with dl2:
    try:
        docx_bytes = generate_docx(report_data)
        st.download_button(
            label="📝 Download Word",
            data=docx_bytes,
            file_name=f"SEO_Weekly_Review_{current_week[0].strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except ImportError:
        st.warning("Install `python-docx` for Word export: `pip install python-docx`")
    except Exception as e:
        st.error(f"Word error: {e}")

with dl3:
    # CSV export of all data
    import io
    csv_buf = io.StringIO()
    export_df = t_weekly[['Week_Label','Sessions','Users','New Users','Active Users']].copy()
    export_df.columns = ['Week','Sessions','Users','New Users','Active Users']
    export_df.to_csv(csv_buf, index=False)
    st.download_button(
        label="📊 Download CSV",
        data=csv_buf.getvalue(),
        file_name=f"SEO_Weekly_Data_{current_week[0].strftime('%Y%m%d')}.csv",
        mime="text/csv",
        use_container_width=True,
    )

# =====================================================
# FOOTER
# =====================================================
st.markdown(f"""
<div style="text-align:center;padding:30px 0 10px;color:#9ca3af;font-size:0.8rem;">
    Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} &nbsp;|&nbsp; Contify SEO Weekly Review
</div>
""", unsafe_allow_html=True)
